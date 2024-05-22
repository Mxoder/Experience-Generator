import os
import docx
import pdfplumber
from loguru import logger
from abc import ABC, abstractmethod
from typing import Optional, Dict, List

from config.config import data_dir

class BaseProcessStrategy(ABC):
    @abstractmethod
    def read_file(self, file_path: str) -> Optional[str]:
        raise NotImplementedError("Subclasses must implement the read_file method.")

class PdfProcessStrategy(BaseProcessStrategy):
    def read_file(self, file_path: str) -> Optional[str]:
        extracted_text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text(
                    x_tolerance=3,
                    x_tolerance_ratio=None,
                    y_tolerance=3,
                    layout=False,
                    x_density=7.25,
                    y_density=13,
                    line_dir_render=None,
                    char_dir_render=None
                )
                extracted_text.append(text)
        return ''.join(extracted_text)
    
class WordProcessStrategy(BaseProcessStrategy):
    def read_file(self, file_path: str) -> Optional[str]:
        doc = docx.Document(file_path)
        text = []

        # Extract text from all paragraphs in the document
        for para in doc.paragraphs:
            text.append(para.text)

        # Extract text from tables in the document
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Using cell.text.strip() here removes any leading or trailing whitespace from the cell text
                    cell_text = cell.text.strip()
                    # Ensure that the cell text is added only if it's not empty
                    if cell_text:
                        text.append(cell_text)
        return '\n'.join(text)
    
class MarkdownProcessStrategy(BaseProcessStrategy):
    def read_file(self, file_path: str) -> Optional[str]:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return text

class TxtProcessStrategy(BaseProcessStrategy):
    def read_file(self, file_path: str) -> Optional[str]:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return text

class FileReader:
    pdf_file_list = []
    word_file_list = []
    md_file_list = []
    txt_file_list = []
    extracted_text = {}

    def __init__(self):
        self._read_files_recursively(data_dir)
        self._extract_text()

    def _read_files_recursively(self, dir_path):
        for root, dirs, files in os.walk(dir_path):
            for file in files:
                file_path = os.path.join(root, file)
                if file.endswith('.pdf'):
                    self.pdf_file_list.append(file_path)
                elif file.endswith(('.doc', '.docx')):
                    self.word_file_list.append(file_path)
                elif file.endswith(('.md', '.markdown')):
                    self.md_file_list.append(file_path)
                elif file.endswith('.txt') or '.' not in file:
                    self.txt_file_list.append(file_path)
                else:
                    logger.warning(f"File type not supported yet: {file}")

    def _extract_text(self):
        strategies = {
            'pdf': PdfProcessStrategy(),
            'word': WordProcessStrategy(),
            'markdown': MarkdownProcessStrategy(),
            'txt': TxtProcessStrategy(),
        }
        
        for file_path in self.pdf_file_list:
            strategy = strategies.get('pdf')
            text = strategy.read_file(file_path)
            title = file_path.split(os.sep)[-1].split('.')[-2].strip()
            self.extracted_text[title] = text

    def get_file_list(self) -> Dict[str, List]:
        return {
            "pdf_file_list": self.pdf_file_list,
            "word_file_list": self.word_file_list,
            "md_file_list": self.md_file_list,
            "txt_file_list": self.txt_file_list,
        }

    def get_extracted_text(self) -> Dict[str, str]:
        return self.extracted_text
